import streamlit as st
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PyPDF2

# Page configuration
st.set_page_config(
    page_title="The TWNC FaloopinFormatter",
    page_icon="🎭",
    layout="wide"
)

st.title("🎭 The TWNC FaloopinFormatter")
st.markdown("Upload your script and we'll reformat it to professional theatrical standards")

# Sidebar
with st.sidebar:
    st.header("📋 How It Works")
    st.markdown("""
    **3 Simple Steps:**
    
    **1️⃣ Upload Your Script**
    - PDF, DOCX, or TXT file
    - We'll read and analyze it
    
    **2️⃣ Enter Details**
    - Provide title and author
    - We'll detect most info automatically
    
    **3️⃣ Download**
    - Choose .DOCX or PDF format
    - Get your professionally formatted script!
    
    ---
    
    **What We Fix:**
    - ✅ Add proper cover page
    - ✅ Add page numbers
    - ✅ Format character names (centered)
    - ✅ Format dialogue (left-aligned)
    - ✅ Format stage directions (italics)
    - ✅ Times New Roman 12pt throughout
    """)

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file and reconstruct paragraphs"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        all_lines = []
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            raw_lines = [line.strip() for line in page_text.split('\n') if line.strip()]
            
            # Reconstruct by merging fragments
            i = 0
            while i < len(raw_lines):
                line = raw_lines[i]
                
                # Skip page numbers
                if line.isdigit():
                    i += 1
                    continue
                
                # Start building a complete line
                complete_line = line
                i += 1
                
                # Keep adding words until we hit a natural break
                while i < len(raw_lines):
                    next_line = raw_lines[i]
                    
                    # Stop if next line is a page number
                    if next_line.isdigit():
                        break
                    
                    # Stop if next line is ALL CAPS (likely character name or heading)
                    if next_line.isupper() and len(next_line.split()) >= 2:
                        break
                    
                    # Stop if current line ends with sentence-ending punctuation
                    if complete_line.rstrip().endswith(('.', '!', '?', ':')):
                        break
                    
                    # Stop if next line is a stage direction label
                    if next_line in ['Yells', 'Softly', 'Laughing', 'Smiling', 'Pauses', 
                                     'Squints', 'Shrugs', 'Opens', 'Looks', 'Takes', 'Reaches',
                                     'Grabs', 'Checks', 'Beat', 'Pause', 'Whistle blows']:
                        break
                    
                    # Add the next fragment
                    if next_line in [',', '.', '!', '?', ':', ';']:
                        complete_line += next_line
                    else:
                        complete_line += " " + next_line
                    i += 1
                    
                    # Stop after reasonable length
                    if len(complete_line) > 200:
                        break
                
                all_lines.append(complete_line)
        
        return "\n".join(all_lines)
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return None

def extract_text_from_txt(txt_file):
    """Extract text from TXT file"""
    try:
        return txt_file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return None

def parse_script_intelligently(text):
    """Parse script text and identify elements"""
    lines = text.split('\n')
    elements = []
    
    # Try to extract title and author from first few lines
    title = None
    author = None
    scene_info = None
    script_start_index = 0
    
    # Look for title in first 10 lines
    for i, line in enumerate(lines[:10]):
        line = line.strip()
        if not line:
            continue
            
        # Check if this might be the title (short, often all caps or title case)
        if not title and len(line) < 50 and not line.startswith('By'):
            title = line
            script_start_index = i + 1
            continue
        
        # Check for author (often has "By" or appears after title)
        if title and not author:
            if line.startswith('By'):
                author = line.replace('By', '').strip()
                script_start_index = i + 1
                continue
            elif i == script_start_index and len(line) < 50:
                author = line
                script_start_index = i + 1
                continue
        
        # Check for scene info
        if re.match(r'(Scene|ACT|Act)', line, re.IGNORECASE):
            scene_info = line
            script_start_index = i + 1
            break
    
    # Now parse the actual script content
    i = script_start_index
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Skip page numbers and headers/footers
        if re.match(r'^\d+$', line):
            i += 1
            continue
        
        # Scene headings
        if re.match(r'^(ACT|SCENE|PROLOGUE|EPILOGUE|INT\.|EXT\.)', line, re.IGNORECASE):
            elements.append({
                'type': 'scene_heading',
                'text': line
            })
            i += 1
            continue
        
        # Setting/stage directions at start
        if re.match(r'^(SETTING:|TIME:|AT RISE|LIGHTS UP|\()', line, re.IGNORECASE):
            elements.append({
                'type': 'setting',
                'text': line
            })
            i += 1
            continue
        
        # Character names - all caps, short line, possibly followed by dialogue
        if line.isupper() and 2 <= len(line.split()) <= 5:
            # Check if next non-empty line exists and isn't all caps
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            
            if next_idx < len(lines):
                next_line = lines[next_idx].strip()
                # If next line is not all caps or is stage direction, this is a character name
                if not next_line.isupper() or next_line.startswith('('):
                    elements.append({
                        'type': 'character',
                        'text': line
                    })
                    i += 1
                    continue
        
        # Standalone stage directions
        if line.startswith('(') and line.endswith(')'):
            # Check if previous element was character name
            if elements and elements[-1]['type'] == 'character':
                # This might be inline with upcoming dialogue
                elements.append({
                    'type': 'dialogue',
                    'text': line
                })
            else:
                elements.append({
                    'type': 'stage_direction',
                    'text': line
                })
            i += 1
            continue
        
        # Stage direction labels (not in parentheses)
        if re.match(r'^[A-Z][a-z]+( [a-z]+)*$', line) and len(line.split()) <= 4:
            # Things like "Exits", "Pause", "Beat"
            if line.lower() in ['beat', 'pause', 'exits', 'enters', 'laughs', 'crying']:
                elements.append({
                    'type': 'stage_direction',
                    'text': f"({line})"
                })
                i += 1
                continue
        
        # Everything else is dialogue
        elements.append({
            'type': 'dialogue',
            'text': line
        })
        i += 1
    
    return title, author, scene_info, elements

def create_formatted_docx(title, author, scene_info, elements):
    """Create professionally formatted Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # COVER PAGE
    for _ in range(8):
        doc.add_paragraph()
    
    # Title
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    # "By"
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("By")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # Scene info at top of first script page
    if scene_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(scene_info)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        doc.add_paragraph()
    
    # Process elements
    for element in elements:
        if element['type'] == 'character':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(12)
            run = p.add_run(element['text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
        elif element['type'] == 'dialogue':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            text = element['text']
            pattern = r'(\([^)]+\))'
            segments = re.split(pattern, text)
            
            for segment in segments:
                if segment:
                    run = p.add_run(segment)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if segment.startswith('(') and segment.endswith(')'):
                        run.italic = True
        
        elif element['type'] == 'stage_direction':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.space_before = Pt(6)
            p.space_after = Pt(6)
            run = p.add_run(element['text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        
        elif element['type'] == 'scene_heading':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(24)
            p.space_after = Pt(12)
            run = p.add_run(element['text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        elif element['type'] == 'setting':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.space_before = Pt(12)
            p.space_after = Pt(12)
            run = p.add_run(element['text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
    
    # END marker
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    run = p.add_run("— END —")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    
    # Page numbers
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio

# Main interface
st.header("🎭 Reformat Your Script in 3 Easy Steps")

# STEP 1: Upload
st.subheader("📤 Step 1: Upload Your Script File")

uploaded_file = st.file_uploader(
    "Choose your script file (PDF, DOCX, or TXT)",
    type=['pdf', 'docx', 'txt'],
    help="Upload a PDF, Word document, or text file"
)

if uploaded_file:
    st.success(f"✅ File uploaded: {uploaded_file.name}")
    
    # Extract text based on file type
    with st.spinner("Reading your script..."):
        if uploaded_file.name.endswith('.pdf'):
            script_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            script_text = extract_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.txt'):
            script_text = extract_text_from_txt(uploaded_file)
        else:
            st.error("Unsupported file type")
            script_text = None
    
    if script_text:
        st.success("✅ Script loaded successfully!")
        
        # Show preview of original
        with st.expander("📄 View Original Text (First 500 characters)"):
            st.text(script_text[:500] + "..." if len(script_text) > 500 else script_text)
        
        # Parse the script
        with st.spinner("Analyzing script structure..."):
            title, author, scene_info, elements = parse_script_intelligently(script_text)
        
        st.info(f"✨ Detected {len(elements)} script elements" + (f" | Scene: {scene_info}" if scene_info else ""))
        
        # STEP 2: Enter title and author
        st.divider()
        st.subheader("✏️ Step 2: Please Enter the Title and Author's Name")
        
        col1, col2 = st.columns(2)
        with col1:
            title = st.text_input("Play Title*", value=title or "", placeholder="e.g., HAMLET")
        with col2:
            author = st.text_input("Author*", value=author or "", placeholder="e.g., William Shakespeare")
        
        # STEP 3: Download
        st.divider()
        st.subheader("📥 Step 3: Download Your Formatted Script")
        
        if not title or not author:
            st.warning("⚠️ Please enter both Title and Author above to continue")
        else:
            filename = st.text_input(
                "Filename (optional - will use title if blank)",
                value="",
                placeholder=title.replace(' ', '_').lower() if title else "my_play"
            )
            
            # Use title as filename if not provided
            if not filename:
                filename = title.replace(' ', '_').lower() if title else "formatted_script"
            
            st.markdown("**Choose your download format:**")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📄 Download as .DOCX", type="primary", use_container_width=True):
                    with st.spinner("Creating Word document..."):
                        formatted_doc = create_formatted_docx(title, author, scene_info, elements)
                    
                    st.success("✅ Document created!")
                    
                    st.download_button(
                        label="💾 Save Word Document",
                        data=formatted_doc,
                        file_name=f"{filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col2:
                if st.button("📕 Download as PDF", type="secondary", use_container_width=True):
                    with st.spinner("Creating PDF document..."):
                        # First create docx
                        formatted_doc = create_formatted_docx(title, author, scene_info, elements)
                        
                        # Try to convert to PDF
                        try:
                            import subprocess
                            import tempfile
                            import os
                            
                            # Save docx temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                                tmp_docx.write(formatted_doc.getvalue())
                                tmp_docx_path = tmp_docx.name
                            
                            # Try to convert with LibreOffice
                            tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
                            result = subprocess.run(
                                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                                 os.path.dirname(tmp_docx_path), tmp_docx_path],
                                capture_output=True,
                                timeout=30
                            )
                            
                            if os.path.exists(tmp_pdf_path):
                                with open(tmp_pdf_path, 'rb') as pdf_file:
                                    pdf_data = pdf_file.read()
                                
                                st.success("✅ PDF created!")
                                
                                st.download_button(
                                    label="💾 Save PDF Document",
                                    data=pdf_data,
                                    file_name=f"{filename}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                                
                                # Cleanup
                                os.unlink(tmp_docx_path)
                                os.unlink(tmp_pdf_path)
                            else:
                                raise Exception("PDF conversion failed")
                        
                        except Exception as e:
                            st.warning("⚠️ PDF conversion not available on this platform")
                            st.info("""
                            **Alternative: Create PDF manually**
                            1. Download the .DOCX file (button on left)
                            2. Open in Microsoft Word or Google Docs
                            3. File → Save As → PDF
                            """)
            
            st.divider()
            st.success("""
            ✨ **Your script will be formatted with:**
            - Professional cover page with title and author
            - Page numbers in upper right corner (starting page 1)
            - Times New Roman 12pt font throughout
            - Character names centered (not bold)
            - Dialogue left-aligned (not indented)
            - Stage directions properly formatted in italics
            - Scene/Act info at top of first script page
            """)

else:
    st.info("👆 Upload a script file above to get started")
    
    st.markdown("""
    ### 📝 Supported Formats:
    - **PDF** (.pdf) - Most common
    - **Word** (.docx) - Microsoft Word documents
    - **Text** (.txt) - Plain text files
    
    ### ✨ What We'll Fix:
    1. Add professional cover page
    2. Add page numbers in upper right corner
    3. Format character names (centered, not bold)
    4. Format dialogue (left-aligned, not indented)
    5. Format stage directions (italics, inline or standalone)
    6. Apply Times New Roman 12pt throughout
    7. Add proper spacing and margins
    """)

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; font-size: 12px; margin-top: 40px;'>
    <p>The TWNC FaloopinFormatter | Professional Script Reformatting</p>
    <p>Upload → Enter Details → Download (.DOCX or PDF)</p>
</div>
""", unsafe_allow_html=True)
